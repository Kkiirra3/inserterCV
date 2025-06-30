from docx import Document
from docxcompose.composer import Composer
from src.services.google_service import GoogleServiceManager
from src.core.template_processor import TemplateProcessor
from src.utils.formatting_utils import FormattingUtils
from src.core.skills_matrix_processor import SkillsMatrixProcessor
from config.config import Config
import os

class DocumentProcessor:
    def __init__(self):
        self.google_service = GoogleServiceManager()
        self.template_processor = TemplateProcessor()
        self.formatting_utils = FormattingUtils()
        self.skills_matrix_processor = SkillsMatrixProcessor()
        
    def merge_google_docs(self, listpage_url, maininfo_url, output_title, template_path=None):
        """
        Main function for merging two Google Docs
        """
        try:
            # Get document IDs
            listpage_id = self.google_service.get_document_id_from_url(listpage_url)
            maininfo_id = self.google_service.get_document_id_from_url(maininfo_url)
            
            # Get services
            drive_service = self.google_service.get_drive_service()
            
            # Create temp directory
            temp_dir = 'temp_docs'
            os.makedirs(temp_dir, exist_ok=True)
            
            # Paths for temporary files
            listpage_docx = os.path.join(temp_dir, 'listpage.docx')
            maininfo_docx = os.path.join(temp_dir, 'maininfo.docx')
            skills_template_docx = os.path.join(temp_dir, 'skills_template.docx')
            projects_template_docx = os.path.join(temp_dir, 'projects_template.docx')
            skills_matrix_template_docx = os.path.join(temp_dir, 'skills_matrix_template.docx')
            skills_matrix_docx = os.path.join(temp_dir, 'skills_matrix.docx')
            merged_docx = os.path.join(temp_dir, 'merged.docx')
            
            # Export documents to .docx
            if not self.google_service.export_to_docx(drive_service, listpage_id, listpage_docx):
                raise Exception("Failed to export LISTPAGE document")
            
            if not self.google_service.export_to_docx(drive_service, maininfo_id, maininfo_docx):
                raise Exception("Failed to export MAIN_INFO document")
            
            # Download skills template
            skills_template_id = "1Xfhp1A7C4OZNxRn1QETSlXR0vj5FcHimJE6TZkQlLJs"
            if not self.google_service.export_to_docx(drive_service, skills_template_id, skills_template_docx):
                raise Exception("Failed to export skills template document")
            
            # Download projects template
            projects_template_id = "1uJUVwNLWG9j_L2HxObvECXhpEAUQ0RRSwTZlJUjh9FA"
            if not self.google_service.export_to_docx(drive_service, projects_template_id, projects_template_docx):
                raise Exception("Failed to export projects template document")

            # Download skills matrix template
            skills_matrix_template_id = Config.INPUT_SKILLS_DOC_ID
            if not self.google_service.export_to_docx(drive_service, skills_matrix_template_id, skills_matrix_template_docx):
                raise Exception("Failed to export skills matrix template document")
            
            # Get formatting from skills template
            skills_doc = Document(skills_template_docx)
            key_para, value_para, key_format, value_format = self.template_processor.find_skills_block_template(skills_doc)
            
            if not (key_format and value_format):
                print("Warning: Could not find formatting in skills template, using default formatting")
            
            if template_path:
                template_data = self.template_processor.load_template_data(template_path)
                
                # Create skills matrix document
                if not self.skills_matrix_processor.create_skills_matrix(
                    skills_matrix_template_docx, 
                    skills_matrix_docx, 
                    template_data
                ):
                    raise Exception("Failed to create skills matrix document")
                
                # First fill projects template with data
                projects_doc = Document(projects_template_docx)
                
                success, bullet_color = self.template_processor.process_projects_template(projects_doc, template_data)
                if not success:
                    raise Exception("Failed to process projects template")
                projects_doc.save(projects_template_docx)
                
                # Process maininfo document
                if not self.template_processor.process_document_with_template(maininfo_docx, template_data, key_format, value_format):
                    raise Exception("Failed to process document with template")
                
                # Find place to insert projects in maininfo
                maininfo_doc = Document(maininfo_docx)
                
                # Remove Tab 1 from main document
                for i, para in enumerate(maininfo_doc.paragraphs):
                    if para.text.strip() == 'Tab 1':
                        p = para._element
                        p.getparent().remove(p)
                        break
                
                # Process projects template marker
                projects_found = False
                # First search in paragraphs
                for para in maininfo_doc.paragraphs:
                    if '{{PROJECTS_TEMPLATE}}' in para.text:
                        projects_found = True
                        
                        # Get parent element
                        parent = para._element.getparent()
                        
                        # Get table from projects_doc
                        for table in projects_doc.tables:
                            # Create deep copy of table
                            table_copy = self.formatting_utils.deepcopy(table._element)
                            # Insert table after paragraph
                            para._element.addnext(table_copy)
                            break
                        
                        # Remove paragraph with marker
                        parent.remove(para._element)
                        break
                
                # If not found in paragraphs, search in tables
                if not projects_found:
                    for table in maininfo_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{PROJECTS_TEMPLATE}}' in para.text:
                                        projects_found = True
                                        
                                        # Get parent element
                                        parent = para._element.getparent()
                                        
                                        # Get table from projects_doc
                                        for proj_table in projects_doc.tables:
                                            # Create deep copy of table
                                            table_copy = self.formatting_utils.deepcopy(proj_table._element)
                                            # Insert table after paragraph
                                            para._element.addnext(table_copy)
                                            break
                                        
                                        # Remove paragraph with marker
                                        parent.remove(para._element)
                                        break
                            if projects_found:
                                break
                        if projects_found:
                            break
                
                if not projects_found:
                    raise Exception("Could not find {{PROJECTS_TEMPLATE}} in main_info document")

                # Process skills matrix marker
                skills_matrix_found = False
                
                # First search in paragraphs
                for para in maininfo_doc.paragraphs:
                    if '{{PROFESSIONAL_SKILLS}}' in para.text:
                        skills_matrix_found = True
                        
                        # Get parent element
                        parent = para._element.getparent()
                        
                        # Get table from skills matrix doc
                        skills_matrix_doc = Document(skills_matrix_docx)
                        for table in skills_matrix_doc.tables:
                            # Create deep copy of table
                            table_copy = self.formatting_utils.deepcopy(table._element)
                            # Insert table after paragraph
                            para._element.addnext(table_copy)
                            break
                        
                        # Remove paragraph with marker
                        parent.remove(para._element)
                        break
                
                # If not found in paragraphs, search in tables
                if not skills_matrix_found:
                    for table in maininfo_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{PROFESSIONAL_SKILLS}}' in para.text:
                                        skills_matrix_found = True
                                        
                                        # Get parent element
                                        parent = para._element.getparent()
                                        
                                        # Get table from skills matrix doc
                                        skills_matrix_doc = Document(skills_matrix_docx)
                                        for matrix_table in skills_matrix_doc.tables:
                                            # Create deep copy of table
                                            table_copy = self.formatting_utils.deepcopy(matrix_table._element)
                                            # Insert table after paragraph
                                            para._element.addnext(table_copy)
                                            break
                                        
                                        # Remove paragraph with marker
                                        parent.remove(para._element)
                                        break
                            if skills_matrix_found:
                                break
                        if skills_matrix_found:
                            break
                
                if not skills_matrix_found:
                    print("Warning: Could not find {{PROFESSIONAL_SKILLS}} in main_info document")
                
                maininfo_doc.save(maininfo_docx)
            
            # Merge .docx files
            if not self.merge_docx_files(listpage_docx, maininfo_docx, merged_docx):
                raise Exception("Failed to merge documents")
            
            # Upload result back to Google Drive with saved bullet points color
            new_doc_id = self.google_service.upload_to_drive(drive_service, merged_docx, output_title, bullet_color)
            if not new_doc_id:
                raise Exception("Failed to upload merged document")
            
            # Form and return URL of new document
            new_doc_url = f"https://docs.google.com/document/d/{new_doc_id}/edit"
            return new_doc_url
        
        except Exception as e:
            print(f"An error occurred: {str(e)}")
            return None

    def merge_docx_files(self, listpage_path, maininfo_path, output_path, template_path=None, key_format=None, value_format=None):
        """
        Merges two .docx files into one using docxcompose
        """
        try:
            # If template path is specified, process only maininfo document
            if template_path:
                template_data = self.template_processor.load_template_data(template_path)
                # Process only maininfo document
                self.template_processor.process_document_with_template(maininfo_path, template_data, key_format, value_format)
            
            # Open base document
            master = Document(listpage_path)
            composer = Composer(master)
            
            # Add second document
            doc2 = Document(maininfo_path)
            composer.append(doc2)
            
            # Save result
            composer.save(output_path)
            return True
            
        except Exception as e:
            print(f"An error occurred while merging: {str(e)}")
            return False 