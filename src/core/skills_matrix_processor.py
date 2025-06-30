from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from datetime import datetime
from typing import List, Dict, Any, Tuple
import re
from config.config import Config


class SkillsMatrixProcessor:
    """Отвечает за все манипуляции с .docx файлом, включая обновление таблицы."""

    def __init__(self):
        self.border_color = Config.BORDER_COLOR
        self.border_size = Config.BORDER_SIZE

    def _set_cell_border(self, cell, **kwargs):
        """Устанавливает или удаляет границы ячейки."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        for edge, edge_data in kwargs.items():
            if edge_data:
                tag = f"w:{edge}"
                elem = OxmlElement(tag)
                for attr, val in edge_data.items():
                    elem.set(qn(f"w:{attr}"), str(val))
                tcBorders.append(elem)

    @staticmethod
    def _cant_split_row(row):
        """Запрещает разрыв строки таблицы между страницами."""
        trPr = row._tr.get_or_add_trPr()
        if not trPr.find(qn('w:cantSplit')):
            trPr.append(OxmlElement('w:cantSplit'))
    
    @staticmethod
    def _change_table_style(table, style_name="TableNormal"):
        """Принудительно меняет стиль таблицы."""
        try:
            tblPr = table._tbl.tblPr
            if tblPr is not None:
                tblStyle = tblPr.find(qn('w:tblStyle'))
                if tblStyle is not None:
                    tblStyle.set(qn('w:val'), style_name)
        except Exception as e:
            print(f"Не удалось изменить стиль таблицы: {e}")

    def _normalize_skill_name(self, tech: str) -> str:
        """Извлекает базовое имя технологии, игнорируя текст в скобках."""
        return tech.split('(')[0].strip()

    def _parse_period(self, period: Dict[str, str]) -> Tuple[int, int]:
        """Парсит период работы над проектом в годы."""
        start = period["start"]
        end = period["end"]
        start_year = int(start.split(".")[1])
        end_year = datetime.now().year if end == "present" else int(end.split(".")[1])
        return start_year, end_year

    def _parse_environment(self, env_list: List[str]) -> List[str]:
        """Разбирает список технологий проекта на отдельные токены."""
        parsed = []
        for item in env_list:
            for token in re.split(r'[(),]', item):
                t = token.strip().lower()
                if t and t != 'etc.':
                    parsed.append(t)
        return parsed

    def _tech_in_env(self, tech: str, env_tokens: List[str]) -> bool:
        """Проверяет, использовалась ли технология в окружении проекта."""
        target_tech = tech.strip().lower()

        if target_tech in env_tokens:
            return True
        if target_tech == "git" and any(token.startswith("git") for token in env_tokens):
            return True
        
        target_parts = target_tech.split()
        if len(target_parts) > 1 and all(part in env_tokens for part in target_parts):
            return True
            
        return False

    def _merge_periods(self, periods: List[Tuple[int, int]]) -> List[Tuple[int, int]]:
        """Объединяет пересекающиеся периоды времени."""
        if not periods:
            return []
        periods.sort()
        merged = [periods[0]]
        for current_start, current_end in periods[1:]:
            last_start, last_end = merged[-1]
            if current_start <= last_end + 1:
                merged[-1] = (last_start, max(last_end, current_end))
            else:
                merged.append((current_start, current_end))
        return merged

    def get_skills_matrix_data(self, template_data: Dict) -> List[List[str]]:
        """Основная функция для анализа и подготовки данных для таблицы."""
        skills_dict = template_data["skills"]["skills"]
        projects = template_data["projects"]

        project_envs = []
        for proj in projects:
            project_envs.append({
                "env": self._parse_environment(proj.get("environment", [])),
                "period": self._parse_period(proj["period"])
            })

        table_data = []
        for category_key, raw_techs in skills_dict.items():
            if category_key == 'introduction':  # Skip introduction
                continue
            category_name = category_key.replace('_', ' ').title()
            category_rows = []
            for tech_full_name in raw_techs:
                norm_name = self._normalize_skill_name(tech_full_name)
                
                periods = [p["period"] for p in project_envs if self._tech_in_env(norm_name, p["env"])]
                merged = self._merge_periods(periods)
                
                if merged:
                    total_exp = sum(end - start + 1 for start, end in merged)
                    last_used = max(end for _, end in merged)
                    category_rows.append([norm_name, str(total_exp), str(last_used)])
                else:
                    category_rows.append([norm_name, "-", "-"])
            
            if category_rows:
                table_data.append([category_name] + category_rows[0])
                for row_data in category_rows[1:]:
                    table_data.append([""] + row_data)

        return table_data

    def update_table(self, data: List[List[str]]):
        """Основная функция для обновления таблицы в документе."""
        if len(self.table.rows) < 2:
            raise RuntimeError("Шаблонная таблица должна содержать минимум 2 строки.")

        self._change_table_style(self.table)

        cat_tmpl = self.table.rows[1].cells[0].paragraphs[0]
        data_tmpls = [c.paragraphs[0] for c in self.table.rows[1].cells[1:]]

        while len(self.table.rows) > 1:
            self.table._tbl.remove(self.table.rows[-1]._tr)

        cat_start_row_idx = -1
        for i, row_vals in enumerate(data):
            is_new_category = bool(row_vals[0])
            new_row = self.table.add_row()
            self._cant_split_row(new_row)

            for idx, cell in enumerate(new_row.cells):
                tmpl = cat_tmpl if idx == 0 else data_tmpls[idx-1]
                text = row_vals[idx] or ''
                cell.text = text
                p = cell.paragraphs[0]

                p.paragraph_format.alignment = tmpl.paragraph_format.alignment
                p.paragraph_format.left_indent = tmpl.paragraph_format.left_indent
                p.paragraph_format.right_indent = tmpl.paragraph_format.right_indent
                p.paragraph_format.space_before = tmpl.paragraph_format.space_before
                p.paragraph_format.space_after = tmpl.paragraph_format.space_after

                if is_new_category:
                    p.paragraph_format.space_before = Pt(12)

                if p.runs and tmpl.runs:
                    rt, run = tmpl.runs[0], p.runs[0]
                    run.font.name, run.font.size, run.font.bold, run.font.italic = rt.font.name, rt.font.size, rt.font.bold, rt.font.italic
                    if rt.font.color and rt.font.color.rgb:
                        run.font.color.rgb = rt.font.color.rgb
                
                self._set_cell_border(cell, top={"val": "nil"}, bottom={"val": "nil"}, left={"val": "nil"}, right={"val": "nil"})

            if is_new_category:
                if cat_start_row_idx != -1:
                    end_row_idx = len(self.table.rows) - 2
                    a, b = self.table.cell(cat_start_row_idx, 0), self.table.cell(end_row_idx, 0)
                    if a != b: a.merge(b)
                    for cell in self.table.rows[end_row_idx].cells:
                        self._set_cell_border(cell, bottom={'sz': self.border_size, 'val': 'single', 'color': self.border_color})
                cat_start_row_idx = len(self.table.rows) - 1

        if cat_start_row_idx != -1:
            end_row_idx = len(self.table.rows) - 1
            a, b = self.table.cell(cat_start_row_idx, 0), self.table.cell(end_row_idx, 0)
            if a != b: a.merge(b)
            for cell in self.table.rows[end_row_idx].cells:
                self._set_cell_border(cell, bottom={'sz': self.border_size, 'val': 'single', 'color': self.border_color})
        

    def create_skills_matrix(self, template_doc_path: str, output_path: str, template_data: Dict) -> bool:
        """Creates skills matrix document based on template"""
        try:
            # Get data for table
            table_data = self.get_skills_matrix_data(template_data)
            
            # Open template document
            doc = Document(template_doc_path)
            if not doc.tables:
                raise RuntimeError("В документе не найдено таблиц.")
            
            self.table = doc.tables[0]
            self.update_table(table_data)
            
            # Save document
            doc.save(output_path)
            return True

        except Exception as e:
            print(f"Error creating skills matrix: {str(e)}")
            return False 