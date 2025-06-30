import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from src.utils.formatting_utils import FormattingUtils
import re
from lxml import etree

class TemplateProcessor:
    def __init__(self):
        self.formatting_utils = FormattingUtils()

    def load_template_data(self, json_path):
        """
        Loads data from template.json
        """
        with open(json_path, 'r') as f:
            return json.load(f)

    def split_introduction(self, intro_text):
        """
        Splits introduction text into two parts by first period
        """
        match = re.search(r'(?<!\d)\.(?!\d)', intro_text)
        if match:
            split_index = match.start()
            first_part = intro_text[:split_index + 1].strip()
            second_part = intro_text[split_index + 1:].strip()
            return first_part, second_part
        return intro_text, ""

    def replace_text_preserve_format(self, paragraph, old_text, new_text):
        """
        Replaces text in paragraph while preserving formatting of each run.
        """
        if not old_text in paragraph.text:
            return False
        
        # Collect all runs and their formatting
        runs_with_text = []
        current_text = ""
        
        for run in paragraph.runs:
            current_text += run.text
            runs_with_text.append({
                'run': run,
                'text': run.text,
                'start': len(current_text) - len(run.text),
                'end': len(current_text),
                'format': {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font': run.font.name if run.font.name else None,
                    'size': run.font.size if run.font.size else None,
                    'color': run.font.color.rgb if run.font.color and run.font.color.rgb else None
                }
            })
        
        # Find old text position
        old_text_start = paragraph.text.index(old_text)
        old_text_end = old_text_start + len(old_text)
        
        # Create new runs with preserved formatting
        new_runs = []
        current_pos = 0
        
        for run_info in runs_with_text:
            if current_pos >= old_text_end:
                # After replacement - copy as is
                new_runs.append((run_info['text'], run_info['format']))
            elif current_pos + len(run_info['text']) <= old_text_start:
                # Before replacement - copy as is
                new_runs.append((run_info['text'], run_info['format']))
            else:
                # During replacement
                if current_pos < old_text_start:
                    # Part before replacement
                    prefix = run_info['text'][:old_text_start - current_pos]
                    if prefix:
                        new_runs.append((prefix, run_info['format']))
                
                # Replaced part
                if current_pos <= old_text_start and current_pos + len(run_info['text']) >= old_text_end:
                    new_runs.append((new_text, run_info['format']))
                
                if current_pos + len(run_info['text']) > old_text_end:
                    # Part after replacement
                    suffix = run_info['text'][old_text_end - current_pos:]
                    if suffix:
                        new_runs.append((suffix, run_info['format']))
            
            current_pos += len(run_info['text'])
        
        # Clear paragraph
        paragraph.clear()
        
        # Add new runs
        for text, format_info in new_runs:
            run = paragraph.add_run(text)
            self.formatting_utils.apply_format_to_run(run, format_info)
        
        return True

    def get_skills_sections(self, template_data):
        """
        Forms lists of keys and values from skills section as separate blocks,
        where each section is added sequentially with the same formatting
        """
        skills = template_data['skills']['skills']
        
        # Format first key-value pair for placeholder replacement
        first_section = True
        result_keys = []
        result_values = []
        first_result = {}
        
        for key, values_list in skills.items():
            # Skip introduction as it's not a skill
            if key == 'introduction':
                continue
                
            # Format header from snake_case to Title Case
            header = key.replace('_', ' ').title()
            
            # Format values
            if isinstance(values_list, list):
                # For lists join elements with comma
                value = ', '.join(str(item) for item in values_list)
            else:
                value = str(values_list)
            
            # Add period at the end if not present
            if not value.endswith('.'):
                value = value + '.'
            
            # For first section use placeholders
            if first_section:
                first_result = {
                    '{{SKILLS_KEY}}': header,
                    '{{SKILLS_VALUE}}': value
                }
                first_section = False
            else:
                # For subsequent sections add to lists
                result_keys.append(header)
                result_values.append(value)
        
        return first_result, result_keys, result_values

    def get_projects_sections(self, template_data):
        """
        Forms list of projects from template.json
        """
        projects = template_data.get('projects', [])
        result = []
        
        for project in projects:
            project_data = {
                '{{PROJECT_NAME}}': project.get('name', ''),
                '{{PROJECT_DESCRIPTION}}': project.get('description', ''),
                '{{PROJECT_ROLES}}': project.get('role', ''),
                '{{PROJECT_PERIOD}}': project.get('period', ''),
                '{{PROJECT_RESPONSIBILITIES}}': project.get('responsibilities', ''),
                '{{PROJECT_ENVIROMENT}}': project.get('environment', '')
            }
            result.append(project_data)
        
        return result

    def find_skills_block_template(self, doc):
        """
        Finds block with placeholders {{SKILLS_KEY}} and {{SKILLS_VALUE}}
        Returns: (key_paragraph, value_paragraph, key_format, value_format)
        """
        key_para = None
        value_para = None
        key_format = None
        value_format = None
        
        for paragraph in doc.paragraphs:
            if '{{SKILLS_KEY}}' in paragraph.text:
                key_para = paragraph
                if paragraph.runs:
                    key_format = {
                        'bold': paragraph.runs[0].bold,
                        'italic': paragraph.runs[0].italic,
                        'underline': paragraph.runs[0].underline,
                        'font': paragraph.runs[0].font.name if paragraph.runs[0].font.name else None,
                        'size': paragraph.runs[0].font.size if paragraph.runs[0].font.size else None,
                        'color': paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color and paragraph.runs[0].font.color.rgb else None,
                        'alignment': paragraph.alignment,
                        'style': paragraph.style
                    }
            elif '{{SKILLS_VALUE}}' in paragraph.text:
                value_para = paragraph
                if paragraph.runs:
                    value_format = {
                        'bold': paragraph.runs[0].bold,
                        'italic': paragraph.runs[0].italic,
                        'underline': paragraph.runs[0].underline,
                        'font': paragraph.runs[0].font.name if paragraph.runs[0].font.name else None,
                        'size': paragraph.runs[0].font.size if paragraph.runs[0].font.size else None,
                        'color': paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color and paragraph.runs[0].font.color.rgb else None,
                        'alignment': paragraph.alignment,
                        'style': paragraph.style
                    }
        
        return key_para, value_para, key_format, value_format

    def find_template_element(self, cell, placeholder):
        """
        Finds template element with specified placeholder and returns its formatting
        """
        for para in cell.paragraphs:
            if placeholder in para.text:
                # Save related document styles
                if hasattr(para, '_element') and hasattr(para._element, 'style'):
                    style_id = para._element.style
                    if style_id:
                        # Copy style definition from document
                        style_element = para._element.getparent().find(f'.//w:style[@w:styleId="{style_id}"]', namespaces=self.formatting_utils.nsmap)
                        if style_element is not None:
                            return deepcopy(para._element), deepcopy(style_element)
                return deepcopy(para._element), None
        return None, None

    def process_projects_template(self, doc, template_data):
        """
        Fills projects template with data from template.json
        """
        projects = template_data.get('projects', [])
        if not projects:
            return False

        # Remove "Tab 1" paragraph
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == 'Tab 1':
                p = para._element
                p.getparent().remove(p)
                break

        # Find table with project placeholders
        template_table = None
        template_row = None
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if '{{PROJECT_NAME}}' in cell.text:
                        template_table = table
                        template_row = row
                        break
                if template_table:
                    break
            if template_table:
                break

        if not template_table or not template_row:
            return False

        # Save formatting for each template element
        first_cell = template_row.cells[0]
        second_cell = template_row.cells[1]
        
        template_formats = {}
        for key, placeholder in [
            ('name', '{{PROJECT_NAME}}'),
            ('description', '{{PROJECT_DESCRIPTION}}'),
            ('roles_header', 'Project roles'),
            ('roles_value', '{{PROJECT_ROLES}}'),
            ('period_header', 'Period'),
            ('period_value', '{{PROJECT_PERIOD}}'),
            ('resp_header', 'Responsibilities'),
            ('resp_value', '{{PROJECT_RESPONSIBILITIES}}'),
            ('env_header', 'Environment'),
            ('env_value', '{{PROJECT_ENVIROMENT}}')
        ]:
            cell = first_cell if key in ['name', 'description'] else second_cell
            element, style = self.find_template_element(cell, placeholder)
            if element is not None:
                # Check if element is a list
                is_list = self.formatting_utils.has_list_properties(element)
                
                # Color will be in rPr of first run
                bullet_color = None
                if key == 'resp_value':
                    first_run = element.find('.//w:r', namespaces=self.formatting_utils.nsmap)
                    if first_run is not None:
                        color_elem = first_run.find('.//w:color', namespaces=self.formatting_utils.nsmap)
                        if color_elem is not None:
                            bullet_color = color_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                
                template_formats[key] = {
                    'element': element,
                    'style': style,
                    'is_list': is_list,
                    'bullet_color': bullet_color
                }
            else:
                template_formats[key] = None

        # Save row template
        template_row_element = deepcopy(template_row._element)

        # Clear table
        while len(template_table.rows) > 0:
            row_element = template_table.rows[0]._element
            row_element.getparent().remove(row_element)

        # Add rows for each project
        for project in projects:
            # Create new row from template
            new_row_element = deepcopy(template_row_element)
            template_table._element.append(new_row_element)
            new_row = template_table.rows[-1]

            # Fill first cell (name and description)
            cell = new_row.cells[0]
            cell._element.clear_content()

            # Project name
            name_para = cell.add_paragraph()
            if template_formats.get('name'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['name']['element'], name_para._element)
                if template_formats['name']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['name']['element'], name_para._element)
                name_run = name_para.add_run(str(project.get('name', '')))
                if template_formats['name']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['name']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], name_run._element)

            # Project description
            desc_para = cell.add_paragraph()
            if template_formats.get('description'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['description']['element'], desc_para._element)
                if template_formats['description']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['description']['element'], desc_para._element)
                desc_run = desc_para.add_run(str(project.get('description', '')))
                if template_formats['description']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['description']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], desc_run._element)

            # Fill second cell (details)
            cell = new_row.cells[1]
            cell._element.clear_content()

            # Roles (header)
            roles_header = cell.add_paragraph()
            if template_formats.get('roles_header'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['roles_header']['element'], roles_header._element)
                if template_formats['roles_header']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['roles_header']['element'], roles_header._element)
                roles_header_run = roles_header.add_run("Project roles")
                if template_formats['roles_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['roles_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], roles_header_run._element)

            # Roles (value)
            roles_value = cell.add_paragraph()
            if template_formats.get('roles_value'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['roles_value']['element'], roles_value._element)
                if template_formats['roles_value']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['roles_value']['element'], roles_value._element)
                roles_run = roles_value.add_run(self.format_value(project.get('role', '')))
                if template_formats['roles_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['roles_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], roles_run._element)

            # Period (header)
            period_header = cell.add_paragraph()
            if template_formats.get('period_header'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['period_header']['element'], period_header._element)
                if template_formats['period_header']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['period_header']['element'], period_header._element)
                period_header_run = period_header.add_run("Period")
                if template_formats['period_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['period_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], period_header_run._element)

            # Period (value)
            period_value = cell.add_paragraph()
            if template_formats.get('period_value'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['period_value']['element'], period_value._element)
                if template_formats['period_value']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['period_value']['element'], period_value._element)
                period_run = period_value.add_run(self.format_value(project.get('period', '')))
                if template_formats['period_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['period_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], period_run._element)

            # Responsibilities (header)
            resp_header = cell.add_paragraph()
            if template_formats.get('resp_header'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['resp_header']['element'], resp_header._element)
                if template_formats['resp_header']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['resp_header']['element'], resp_header._element)
                resp_header_run = resp_header.add_run("Responsibilities")
                if template_formats['resp_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['resp_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], resp_header_run._element)

            # Responsibilities (values)
            responsibilities = project.get('responsibilities', [])
            if isinstance(responsibilities, list):
                for i, resp in enumerate(responsibilities):
                    resp_value = cell.add_paragraph()
                    if template_formats.get('resp_value'):
                        # Copy paragraph and list formatting
                        self.formatting_utils.copy_paragraph_format_with_ns(template_formats['resp_value']['element'], resp_value._element)
                        if template_formats['resp_value']['is_list']:
                            self.formatting_utils.copy_list_properties(template_formats['resp_value']['element'], resp_value._element)
                        # Add semicolon for all items except last one, which gets a period
                        resp_text = str(resp) + (';' if i < len(responsibilities) - 1 else '.')
                        resp_run = resp_value.add_run(resp_text)
                        if template_formats['resp_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                            self.formatting_utils.copy_run_format_with_ns(template_formats['resp_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], resp_run._element)

            # Environment (header)
            env_header = cell.add_paragraph()
            if template_formats.get('env_header'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['env_header']['element'], env_header._element)
                if template_formats['env_header']['is_list']:
                    self.formatting_utils.copy_list_properties(template_formats['env_header']['element'], env_header._element)
                env_header_run = env_header.add_run("Environment")
                if template_formats['env_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['env_header']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], env_header_run._element)

            # Environment (value) - format as comma-separated string and add period at end
            env_value = cell.add_paragraph()
            if template_formats.get('env_value'):
                self.formatting_utils.copy_paragraph_format_with_ns(template_formats['env_value']['element'], env_value._element)
                env_text = self.format_value(project.get('environment', ''))
                if not env_text.endswith('.'):
                    env_text += '.'
                env_run = env_value.add_run(env_text)
                if template_formats['env_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap):
                    self.formatting_utils.copy_run_format_with_ns(template_formats['env_value']['element'].findall('.//w:r', namespaces=self.formatting_utils.nsmap)[0], env_run._element)

        return True, template_formats.get('resp_value', {}).get('bullet_color')

    def format_value(self, value):
        """
        Formats value for document insertion
        """
        if isinstance(value, dict):
            if 'start' in value and 'end' in value:
                return f"{value['start']} - {value['end']}"
        elif isinstance(value, list):
            # For environment join values with comma without spaces
            return ', '.join(str(item) for item in value)
        return str(value)

    def process_document_with_template(self, doc_path, template_data, key_format=None, value_format=None):
        """
        Processes document, replacing placeholders with template data
        """
        doc = Document(doc_path)
        
        # Get introduction parts
        intro_part1, intro_part2 = self.split_introduction(template_data['skills']['introduction'])
        
        # Get basic information
        basic_info = template_data['skills']['basic_information']
        
        # Dictionary of basic replacements with exact placeholders
        replacements = {
            '{{NAME}}': template_data['personal_info']['name'],
            '{{TITLE}}': template_data['personal_info']['title'],
            '{{EDUCATION_TEMPLATE}}': basic_info['education'],
            '{{LANGUAGES}}': self.format_skills_list(basic_info['languages']),
            '{{DOMAINS_TEMPLATE}}': self.format_domains_list(basic_info['domains'])
        }
        
        # Process basic placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    self.replace_text_preserve_format(paragraph, placeholder, str(value))
        
        # Process placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    # First process INTRO blocks
                    intro_para = None
                    intro_format = None
                    
                    # Find paragraph with INTRO_PART_1
                    for para in cell.paragraphs:
                        if '{{INTRO_PART_1}}' in para.text:
                            intro_para = para
                            if para.runs:
                                intro_format = {
                                    'bold': para.runs[0].bold,
                                    'italic': para.runs[0].italic,
                                    'underline': para.runs[0].underline,
                                    'font': para.runs[0].font.name if para.runs[0].font.name else None,
                                    'size': para.runs[0].font.size if para.runs[0].font.size else None,
                                    'color': para.runs[0].font.color.rgb if para.runs[0].font.color and para.runs[0].font.color.rgb else None,
                                    'alignment': para.alignment,
                                    'style': para.style,
                                    'paragraph_format': para.paragraph_format
                                }
                            break
                    
                    # If found INTRO_PART_1, process it and INTRO_PART_2
                    if intro_para and intro_format:
                        # Replace INTRO_PART_1
                        self.replace_text_preserve_format(intro_para, '{{INTRO_PART_1}}', intro_part1)
                        
                        # Find paragraph with INTRO_PART_2 and process it separately
                        for para in cell.paragraphs:
                            if '{{INTRO_PART_2}}' in para.text:
                                # Save paragraph formatting
                                if intro_format.get('paragraph_format'):
                                    if hasattr(intro_format['paragraph_format'], 'left_indent'):
                                        para.paragraph_format.left_indent = intro_format['paragraph_format'].left_indent
                                    if hasattr(intro_format['paragraph_format'], 'first_line_indent'):
                                        para.paragraph_format.first_line_indent = intro_format['paragraph_format'].first_line_indent
                                
                                # Replace text preserving formatting
                                self.replace_text_preserve_format(para, '{{INTRO_PART_2}}', intro_part2)
                                
                                # Apply exact same formatting as INTRO_PART_1
                                for run in para.runs:
                                    run.bold = False  # Force remove bold
                                    if intro_format.get('font'):
                                        run.font.name = intro_format['font']
                                    if intro_format.get('size'):
                                        run.font.size = intro_format['size']
                                    if intro_format.get('color'):
                                        if not run.font.color:
                                            run.font.color = RGBColor(0, 0, 0)  # Initialize color if none exists
                                        run.font.color.rgb = intro_format['color']
                    
                    # Process other placeholders
                    for para in cell.paragraphs:
                        for placeholder, value in replacements.items():
                            if placeholder in para.text:
                                self.replace_text_preserve_format(para, placeholder, str(value))
        
        # If this is main_info document, process skills sections
        if 'maininfo' in doc_path.lower():
            # Process skills section
            skills_table = None
            skills_cell = None
            target_para = None
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '{{SKILLS_FABRYC}}' in para.text:
                                skills_table = table
                                skills_cell = cell
                                target_para = para
                                break
                        if skills_table:
                            break
                    if skills_table:
                        break
                if skills_table:
                    break
            
            if skills_table and target_para:
                # Clear cell content, saving only paragraphs before marker
                paras_before = []
                for para in skills_cell.paragraphs:
                    if '{{SKILLS_FABRYC}}' in para.text:
                        break
                    paras_before.append(para._p)
                
                # Clear cell
                skills_cell._element.clear_content()
                
                # Restore paragraphs before marker
                for p in paras_before:
                    skills_cell._element.append(p)
                
                # Insert all skills sections
                skills = template_data['skills']['skills']
                
                for key, values_list in skills.items():
                    if key == 'introduction':
                        continue
                    
                    try:
                        # Format header
                        header = key.replace('_', ' ').title()
                        
                        # Format values
                        if isinstance(values_list, list):
                            value = ', '.join(str(item) for item in values_list)
                        else:
                            value = str(values_list)
                        
                        # Add period at the end if not present
                        if not value.endswith('.'):
                            value = value + '.'
                        
                        # Create paragraph for key and value
                        para = skills_cell.add_paragraph()
                        if key_format:
                            # Copy paragraph formatting
                            if key_format.get('style'):
                                para.style = key_format['style']
                            if key_format.get('alignment'):
                                para.alignment = key_format['alignment']
                            if key_format.get('paragraph_format'):
                                if hasattr(key_format['paragraph_format'], 'left_indent'):
                                    para.paragraph_format.left_indent = key_format['paragraph_format'].left_indent
                                if hasattr(key_format['paragraph_format'], 'first_line_indent'):
                                    para.paragraph_format.first_line_indent = key_format['paragraph_format'].first_line_indent
                        
                        # Add key (header)
                        key_run = para.add_run(header)
                        if key_format:
                            key_run.bold = True  # Headers always bold
                            if key_format.get('font'):
                                key_run.font.name = key_format['font']
                            if key_format.get('size'):
                                key_run.font.size = key_format['size']
                            if key_format.get('color'):
                                key_run.font.color.rgb = key_format['color']
                        
                        # Add line break
                        para.add_run('\n')
                        
                        # Add value in same paragraph
                        value_run = para.add_run(value)
                        if value_format:
                            value_run.bold = False  # Values should not be bold
                            if value_format.get('font'):
                                value_run.font.name = value_format['font']
                            if value_format.get('size'):
                                value_run.font.size = value_format['size']
                            if value_format.get('color'):
                                value_run.font.color.rgb = value_format['color']
                        
                        # Set small spacing between blocks and allow page breaks
                        para.paragraph_format.space_before = Pt(0)  # No space before
                        para.paragraph_format.space_after = Pt(9)  # Small space after
                        para.paragraph_format.keep_together = True  # Allow page breaks
                        para.paragraph_format.keep_with_next = False  # Don't keep with next paragraph
                    
                    except Exception as e:
                        print(f"Error adding section {key}: {str(e)}")
                        continue
        
        doc.save(doc_path)
        return True

    def format_skills_list(self, skills_list):
        """
        Formats skills list as string
        """
        if isinstance(skills_list, list):
            return ', '.join(str(item) for item in skills_list)
        return str(skills_list) 

    def format_domains_list(self, skills_list):
        """
        Formats domains list as string
        """
        if isinstance(skills_list, list):
            return ', \n'.join(str(item) for item in skills_list)
        return str(skills_list) 